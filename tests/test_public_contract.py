from __future__ import annotations

import base64
import sys
import tempfile
import types
import zipfile
from pathlib import Path
from typing import Any, cast
from unittest import mock

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from docxrender import (
    DocxBodyAnchorOptions,
    DocxBodyRenderPolicy,
    DocxFieldMarkerOptions,
    DocxFieldRefreshOptions,
    DocxFontStyle,
    DocxHeaderFooterImageOptions,
    DocxMarkdownOptions,
    DocxParagraphStyle,
    DocxRenderer,
    DocxSizeStyle,
    DocxStyle,
    DocxTableStyle,
    DocxTemplateContextPolicy,
    DocxTemplateImageSpec,
    DocxTemplateRenderOptions,
    DocxTemplateRenderResult,
    DocxToPdfOptions,
    DocxToPdfResult,
    DocxWriteOptions,
    DocxWriteResult,
    convert_docx_to_pdf,
    write_docx,
    write_docx_template,
)
from docxrender.docx.fields import (
    write_docx_field_update_markers,
    write_frozen_docx_fields,
)
from docxrender.markdown import (
    MarkdownHeading,
    MarkdownImage,
    MarkdownOrderedList,
    MarkdownParagraph,
    MarkdownTable,
    MarkdownTextSegment,
    MarkdownUnorderedList,
    parse_markdown_blocks,
)
from docxrender.pdf_uno import (
    create_libreoffice_listener_command,
    create_load_failure_fields,
    import_uno_module,
    start_libreoffice_listener,
    wait_for_listener,
)


class FakeProcess:
    def __init__(self) -> None:
        self.terminated = False

    def poll(self) -> int:
        return 0

    def terminate(self) -> None:
        self.terminated = True

    def wait(self, timeout: float | None = None) -> int:
        return 0

    def kill(self) -> None:
        self.terminated = True


class FakeDocument:
    def __init__(self) -> None:
        self.closed = False
        self.stored = False
        self.store_url_calls: list[tuple[str, tuple[Any, ...]]] = []

    def store(self) -> None:
        self.stored = True

    def storeToURL(self, url: str, properties: tuple[Any, ...]) -> None:
        self.store_url_calls.append((url, properties))

    def close(self, deliver_ownership: bool) -> None:
        self.closed = True


def create_fake_file_url(path: str) -> str:
    return f"file://{path}"


def create_fake_property(name: str, value: object) -> tuple[str, object]:
    return name, value


def create_docx_style() -> DocxStyle:
    return DocxStyle(
        fonts=DocxFontStyle(
            font_name_latin="Times New Roman",
            font_name_body_east_asia="宋体",
            font_name_heading_east_asia="宋体",
        ),
        sizes=DocxSizeStyle(
            pt_title_page_title=36.0,
            pt_title_page_meta=18.0,
            pt_title_page_compiler=15.0,
            pt_body=12.0,
            pt_caption=10.5,
            pt_table=12.0,
            pt_heading_by_level={1: 16.0, 2: 14.0, 3: 12.0},
        ),
        table=DocxTableStyle(
            border_color="000000",
            stripe_fill_color="D9D9D9",
            border_size_main="12",
            border_size_header="6",
            line_spacing=1.5,
        ),
        paragraph=DocxParagraphStyle(
            line_spacing_body=1.5,
            line_spacing_note=1.2,
            first_line_indent_cm=0.74,
        ),
    )


class TestPublicContract:
    def test_public_imports_are_explicit(self) -> None:
        import docxrender

        assert docxrender.__all__ == [
            "DocxRenderer",
            "DocxBodyAnchorOptions",
            "DocxBodyRenderPolicy",
            "DocxFieldMarkerOptions",
            "DocxFieldRefreshOptions",
            "DocxFontStyle",
            "DocxHeaderFooterImageOptions",
            "DocxMarkdownOptions",
            "DocxParagraphStyle",
            "DocxSizeStyle",
            "DocxStyle",
            "DocxTableStyle",
            "DocxTemplateContextPolicy",
            "DocxTemplateImageSpec",
            "DocxTemplateRenderOptions",
            "DocxTemplateRenderResult",
            "DocxToPdfOptions",
            "DocxToPdfResult",
            "DocxWriteOptions",
            "DocxWriteResult",
            "convert_docx_to_pdf",
            "write_docx_template",
            "write_docx",
        ]

    def test_docx_write_options_construct_from_structured_inputs(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            options = DocxWriteOptions(
                file_template=path_tmp / "template.docx",
                file_out_docx=path_tmp / "report.docx",
                context={"report_title": "Example"},
                markdown_body="# Heading\n\nBody.",
                dir_base=path_tmp,
                style=create_docx_style(),
            )

            assert options.body_anchor.anchor_token == "__REPORT_BODY_ANCHOR__"
            assert options.body_anchor.rule_match == "equals"
            assert options.body_anchor.rule_missing == "append"
            assert options.markdown.should_parse_inline_bold is True
            assert options.markdown.should_parse_inline_code is True
            assert options.markdown.should_parse_links_as_text is True
            assert options.markdown.should_parse_image_width_attr is True
            assert options.markdown.default_image_width_pct == 90.0
            assert options.should_update_fields is True
            assert options.should_freeze_fields is False
            assert options.field_refresh is None
            assert options.header_footer_images is None
            assert options.style.paragraph.first_line_indent_cm == 0.74

    def test_docx_body_anchor_options_default_to_equals_append(self) -> None:
        options = DocxBodyAnchorOptions()

        assert options.anchor_token == "__REPORT_BODY_ANCHOR__"
        assert options.rule_match == "equals"
        assert options.rule_missing == "append"
        assert options.should_preserve_section_properties is True

    def test_docx_to_pdf_options_construct_from_conversion_inputs(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            options = DocxToPdfOptions(
                exe_libreoffice=Path("/usr/bin/libreoffice"),
                file_in_docx=path_tmp / "report.docx",
                file_out_pdf=path_tmp / "report.pdf",
                dir_user_profile=path_tmp / "lo-profile",
            )

            assert options.file_out_docx_refreshed is None
            assert options.file_listener_log is None
            assert options.should_update_fields is True
            assert options.should_freeze_fields is False

    def test_docx_markdown_options_default_to_commonmarkish_subset(self) -> None:
        options = DocxMarkdownOptions()

        assert options.should_parse_inline_bold is True
        assert options.should_parse_inline_code is True
        assert options.should_parse_links_as_text is True
        assert options.should_parse_image_width_attr is True
        assert options.default_image_width_pct == 90.0

    def test_docx_body_render_policy_defaults_to_current_renderer_behavior(
        self,
    ) -> None:
        policy = DocxBodyRenderPolicy()

        assert policy.should_number_headings is False
        assert policy.rule_ordered_list == "word_style"
        assert policy.rule_unordered_list == "word_style"
        assert policy.should_stripe_table_rows is False

    def test_docx_template_render_options_default_context_defaults_to_empty(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            options = DocxTemplateRenderOptions(
                file_template=path_tmp / "template.docx",
                file_out_docx=path_tmp / "report.docx",
                context={"report_title": "Example"},
            )

            assert options.context_defaults == {}
            assert options.context_policy == DocxTemplateContextPolicy()

    def test_docx_template_context_policy_defaults_to_caller_wins(self) -> None:
        policy = DocxTemplateContextPolicy()

        assert policy.rule_merge == "merge"
        assert policy.rule_conflict == "caller_wins"
        assert policy.required_keys == ()

    def test_public_results_are_structured_paths(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            result_template = DocxTemplateRenderResult(
                file_docx=path_tmp / "template-rendered.docx"
            )
            result_docx = DocxWriteResult(file_docx=path_tmp / "report.docx")
            result_pdf = DocxToPdfResult(file_pdf=path_tmp / "report.pdf")

            assert result_template.file_docx.name == "template-rendered.docx"
            assert result_docx.file_docx.name == "report.docx"
            assert result_pdf.file_pdf.name == "report.pdf"
            assert result_pdf.file_docx_refreshed is None

    def test_docx_size_style_with_overrides_changes_selected_values(self) -> None:
        sizes = create_docx_style().sizes

        updated = sizes.with_overrides(pt_body=11.0)

        assert updated.pt_body == 11.0
        assert updated.pt_caption == sizes.pt_caption
        assert updated.pt_heading_by_level == sizes.pt_heading_by_level

    def test_docx_size_style_with_overrides_copies_heading_sizes(self) -> None:
        sizes = create_docx_style().sizes
        heading_sizes = {1: 18.0, 2: 15.0}

        updated = sizes.with_overrides(pt_heading_by_level=heading_sizes)
        heading_sizes[1] = 99.0

        assert updated.pt_heading_by_level == {1: 18.0, 2: 15.0}
        assert sizes.pt_heading_by_level[1] == 16.0

    def test_docx_style_component_overrides_change_selected_values(self) -> None:
        style = create_docx_style()

        updated = style.with_overrides(
            fonts=style.fonts.with_overrides(font_name_body_east_asia="黑体"),
            table=style.table.with_overrides(stripe_fill_color="FFFFFF"),
            paragraph=style.paragraph.with_overrides(note_prefixes=("Note:",)),
        )

        assert updated.fonts.font_name_latin == "Times New Roman"
        assert updated.fonts.font_name_body_east_asia == "黑体"
        assert updated.table.border_color == "000000"
        assert updated.table.stripe_fill_color == "FFFFFF"
        assert updated.paragraph.line_spacing_body == 1.5
        assert updated.paragraph.note_prefixes == ("Note:",)

    def test_docx_field_marker_options_default_to_update_only(self) -> None:
        options = DocxFieldMarkerOptions()

        assert options.should_update_fields is True
        assert options.should_freeze_fields is False

    def test_docx_renderer_style_returns_default_style(self) -> None:
        style = DocxRenderer().style

        assert style.fonts.font_name_latin == "Times New Roman"
        assert style.fonts.font_name_body_east_asia == "宋体"
        assert style.sizes.pt_body == 12.0
        assert style.sizes.pt_heading_by_level[1] == 16.0
        assert style.sizes.pt_heading_by_level[6] == 12.0
        assert style.table.border_color == "000000"
        assert style.paragraph.first_line_indent_cm == 0.74

    def test_docx_renderer_fluent_overrides_are_partial(self) -> None:
        style = (
            DocxRenderer()
            .with_fonts(font_name_body_east_asia="黑体")
            .with_sizes(pt_body=11.0)
            .with_table(stripe_fill_color="FFFFFF")
            .with_paragraph(note_prefixes=("Note:",))
            .style
        )

        assert style.fonts.font_name_latin == "Times New Roman"
        assert style.fonts.font_name_body_east_asia == "黑体"
        assert style.sizes.pt_body == 11.0
        assert style.sizes.pt_caption == 10.5
        assert style.table.border_color == "000000"
        assert style.table.stripe_fill_color == "FFFFFF"
        assert style.paragraph.note_prefixes == ("Note:",)
        assert style.paragraph.line_spacing_body == 1.5

    def test_docx_renderer_with_markdown_overrides_markdown_options(self) -> None:
        renderer = DocxRenderer().with_markdown(
            should_parse_inline_bold=False,
            should_parse_inline_code=False,
            should_parse_links_as_text=False,
            should_parse_image_width_attr=False,
            default_image_width_pct=75.0,
        )

        assert renderer.markdown.should_parse_inline_bold is False
        assert renderer.markdown.should_parse_inline_code is False
        assert renderer.markdown.should_parse_links_as_text is False
        assert renderer.markdown.should_parse_image_width_attr is False
        assert renderer.markdown.default_image_width_pct == 75.0

    def test_docx_renderer_with_body_render_policy_overrides_policy(self) -> None:
        renderer = DocxRenderer().with_body_render_policy(
            should_number_headings=True,
            rule_ordered_list="plain_text",
            rule_unordered_list="plain_text",
            should_stripe_table_rows=True,
        )

        assert renderer.body_render_policy == DocxBodyRenderPolicy(
            should_number_headings=True,
            rule_ordered_list="plain_text",
            rule_unordered_list="plain_text",
            should_stripe_table_rows=True,
        )

    def test_docx_renderer_with_template_copies_context(self) -> None:
        context = {"report_title": "Example"}

        renderer = DocxRenderer().with_template(context=context)
        context["report_title"] = "Changed"

        assert renderer.template_context == {"report_title": "Example"}

    def test_docx_renderer_with_template_returns_new_renderer(self) -> None:
        base = DocxRenderer()
        derived = base.with_template(
            file_template=Path("template.docx"),
            context={"report_title": "Example"},
        )

        assert base is not derived
        assert base.template_context == {}
        assert base.file_docx is None
        assert derived.template_context == {"report_title": "Example"}

    def test_docx_renderer_with_template_sets_policy(self) -> None:
        renderer = DocxRenderer().with_template(
            rule_conflict="defaults_win",
            required_keys=("report_title", "body_anchor"),
        )

        assert renderer.template_context_policy == DocxTemplateContextPolicy(
            rule_merge="merge",
            rule_conflict="defaults_win",
            required_keys=("report_title", "body_anchor"),
        )

    def test_docx_renderer_with_template_copies_inline_images(self) -> None:
        inline_images = {
            "cover_image": DocxTemplateImageSpec(
                file_image=Path("cover.png"),
                width_mm=80,
            )
        }

        renderer = DocxRenderer().with_template(inline_images=inline_images)
        inline_images["cover_image"] = DocxTemplateImageSpec(
            file_image=Path("other.png"),
            width_mm=40,
        )

        assert renderer.build_options(
            file_template=Path("template.docx"),
            file_out_docx=Path("report.docx"),
            context={},
            markdown_body="Body.",
            dir_base=Path("."),
        ).template_inline_images == {
            "cover_image": DocxTemplateImageSpec(
                file_image=Path("cover.png"),
                width_mm=80,
            )
        }

    def test_docx_renderer_build_style_matches_style_property(self) -> None:
        renderer = DocxRenderer().with_sizes(pt_body=11.0)

        assert renderer.build_style() is renderer.style

    def test_docx_renderer_build_options_uses_fluent_state(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            field_refresh = DocxFieldRefreshOptions(
                exe_libreoffice=Path("/usr/bin/libreoffice"),
                dir_user_profile=path_tmp / "lo-profile",
            )
            header_footer = DocxHeaderFooterImageOptions(
                file_header_image=path_tmp / "header.png",
                idx_section_start=1,
            )

            options = (
                DocxRenderer()
                .with_template(
                    file_template=path_tmp / "template.docx",
                    context={"report_title": "Builder"},
                )
                .with_sizes(pt_body=11.0)
                .with_field_update_markers(should_update_fields=False)
                .with_field_refresh(field_refresh)
                .with_header_footer_images(header_footer)
                .build_options(
                    file_out_docx=path_tmp / "report.docx",
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            assert options.style.sizes.pt_body == 11.0
            assert options.should_update_fields is False
            assert options.should_freeze_fields is False
            assert options.field_refresh is field_refresh
            assert options.header_footer_images is header_footer
            header_footer_options = options.header_footer_images
            assert header_footer_options is not None
            assert header_footer_options.idx_section_start == 1
            assert options.body_anchor.anchor_token == "__REPORT_BODY_ANCHOR__"
            renderer = DocxRenderer().with_field_refresh(field_refresh)
            built = renderer.build_options(
                file_template=path_tmp / "template.docx",
                file_out_docx=path_tmp / "report.docx",
                context={},
                markdown_body="Body.",
                dir_base=path_tmp,
            )
            assert renderer.docx_options is None
            assert built.field_refresh is field_refresh

    def test_docx_renderer_build_options_carries_body_render_policy(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            policy = DocxBodyRenderPolicy(
                should_number_headings=True,
                rule_ordered_list="plain_text",
            )
            built = (
                DocxRenderer()
                .with_body_render_policy(policy)
                .build_options(
                    file_template=path_tmp / "template.docx",
                    file_out_docx=path_tmp / "report.docx",
                    context={},
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            assert built.body_render_policy is policy

    def test_docx_renderer_build_options_carries_markdown_options(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            built = (
                DocxRenderer()
                .with_template(
                    file_template=path_tmp / "template.docx",
                    context={"report_title": "Builder"},
                )
                .with_markdown(default_image_width_pct=72.0)
                .build_options(
                    file_out_docx=path_tmp / "report.docx",
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            assert built.markdown.default_image_width_pct == 72.0

    def test_docx_renderer_build_options_can_use_stored_template_context(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            built = (
                DocxRenderer()
                .with_template(
                    file_template=path_tmp / "template.docx",
                    context={"report_title": "Stored"},
                )
                .build_options(
                    file_out_docx=path_tmp / "report.docx",
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            assert built.context == {"report_title": "Stored"}
            assert built.template_context_policy == DocxTemplateContextPolicy()

    def test_docx_renderer_build_options_carries_template_context_policy(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            built = (
                DocxRenderer()
                .with_template(
                    file_template=path_tmp / "template.docx",
                    context={"report_title": "Stored"},
                    required_keys=("report_title",),
                )
                .build_options(
                    file_out_docx=path_tmp / "report.docx",
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            assert built.template_context_policy == DocxTemplateContextPolicy(
                required_keys=("report_title",)
            )

    def test_docx_renderer_body_anchor_can_be_built_from_keywords(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            options = (
                DocxRenderer()
                .with_body_anchor(
                    anchor_token="BODY-HERE",
                    rule_match="contains",
                    rule_missing="raise",
                    should_preserve_section_properties=False,
                )
                .build_options(
                    file_template=path_tmp / "template.docx",
                    file_out_docx=path_tmp / "report.docx",
                    context={},
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            assert options.body_anchor == DocxBodyAnchorOptions(
                anchor_token="BODY-HERE",
                rule_match="contains",
                rule_missing="raise",
                should_preserve_section_properties=False,
            )

    def test_docx_renderer_header_footer_images_can_be_built_from_keywords(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            options = (
                DocxRenderer()
                .with_header_footer_images(
                    file_header_image=path_tmp / "header.png",
                    idx_section_start=1,
                )
                .build_options(
                    file_template=path_tmp / "template.docx",
                    file_out_docx=path_tmp / "report.docx",
                    context={},
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            assert options.header_footer_images is not None
            assert options.header_footer_images.idx_section_start == 1

    def test_docx_renderer_field_refresh_can_be_built_from_keywords(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)

            options = (
                DocxRenderer()
                .with_field_refresh(
                    exe_libreoffice=Path("/usr/bin/libreoffice"),
                    dir_user_profile=path_tmp / "lo-profile",
                    should_require_toc=True,
                )
                .build_options(
                    file_template=path_tmp / "template.docx",
                    file_out_docx=path_tmp / "report.docx",
                    context={},
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            assert options.field_refresh is not None
            field_refresh = options.field_refresh
            assert field_refresh is not None
            assert field_refresh.exe_libreoffice == Path("/usr/bin/libreoffice")
            assert field_refresh.should_require_toc is True

    def test_docx_renderer_requires_field_refresh_runtime_paths(self) -> None:
        with pytest.raises(ValueError, match="exe_libreoffice"):
            DocxRenderer().with_field_refresh()

    def test_docx_renderer_field_markers_can_be_built_from_options(self) -> None:
        marker_options = DocxFieldMarkerOptions(
            should_update_fields=False,
            should_freeze_fields=True,
        )
        renderer = DocxRenderer().with_field_update_markers(marker_options)

        assert renderer.field_markers is marker_options

    def test_docx_renderer_build_options_requires_template_file(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)

            with pytest.raises(ValueError, match="file_template is required"):
                DocxRenderer().build_options(
                    file_out_docx=path_tmp / "report.docx",
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )

    def test_docx_writer_is_not_public(self) -> None:
        import docxrender

        assert not hasattr(docxrender, "DocxWriter")

    def test_docx_renderer_write_docx_uses_core_writer(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template(file_template)

            result = (
                DocxRenderer()
                .with_sizes(pt_body=11.0)
                .write_docx(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={"report_title": "Builder Report"},
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            document = Document(str(result.file_docx))
            paragraph_by_text = {
                paragraph.text: paragraph for paragraph in document.paragraphs
            }
            assert result.file_docx == file_out_docx
            assert (
                _run_font_size_pt(_first_text_run(paragraph_by_text["Body."])) == 11.0
            )

    def test_docx_renderer_can_write_template_only(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "rendered.docx"
            _write_template_with_anchor_text(file_template, "{{ body_anchor }}")

            result = (
                DocxRenderer()
                .with_template(
                    file_template=file_template,
                    context={"report_title": "Template Only"},
                )
                .write_docx_template(
                    file_out_docx=file_out_docx,
                )
            )

            texts = [
                paragraph.text
                for paragraph in Document(str(result.file_docx)).paragraphs
            ]
            assert result.file_docx == file_out_docx
            assert "Template Only" in texts
            assert "{{ body_anchor }}" not in texts

    def test_docx_renderer_write_docx_can_use_stored_template_context(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template(file_template)

            result = (
                DocxRenderer()
                .with_template(
                    file_template=file_template,
                    context={"report_title": "Stored Context"},
                )
                .write_docx(
                    file_out_docx=file_out_docx,
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )
            )

            texts = [
                paragraph.text
                for paragraph in Document(str(result.file_docx)).paragraphs
            ]
            assert result.file_docx == file_out_docx
            assert "Stored Context" in texts

    def test_docx_renderer_clone_copies_state_without_linking(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            base = (
                DocxRenderer(file_docx=path_tmp / "report.docx")
                .with_sizes(pt_body=11.0)
                .with_pdf_conversion(
                    exe_libreoffice=Path("/usr/bin/libreoffice"),
                    dir_user_profile=path_tmp / "lo-profile",
                    file_out_pdf=path_tmp / "report.pdf",
                )
            )
            clone = DocxRenderer(base, file_docx=path_tmp / "edited.docx").with_sizes(
                pt_body=10.0
            )

            assert base.file_docx == path_tmp / "report.docx"
            assert clone.file_docx == path_tmp / "edited.docx"
            assert base.style.sizes.pt_body == 11.0
            assert clone.style.sizes.pt_body == 10.0
            assert clone.pdf_options is not None
            assert clone.pdf_options.file_in_docx == path_tmp / "edited.docx"

    def test_docx_renderer_can_postprocess_existing_docx(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_docx = path_tmp / "report.docx"
            _write_template(file_docx)

            result = DocxRenderer(file_docx=file_docx).write_docx()

            assert result.file_docx == file_docx
            with zipfile.ZipFile(file_docx) as zip_file:
                settings = zip_file.read("word/settings.xml").decode("utf-8")
            assert "<w:updateFields" in settings

    def test_docx_renderer_can_disable_existing_docx_field_markers(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_docx = path_tmp / "report.docx"
            _write_template(file_docx)

            DocxRenderer(file_docx=file_docx).with_field_update_markers(
                should_update_fields=False,
            ).write_docx()

            with zipfile.ZipFile(file_docx) as zip_file:
                settings = zip_file.read("word/settings.xml").decode("utf-8")
            assert "<w:updateFields" not in settings

    def test_docx_renderer_write_pdf_uses_current_docx(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_docx = path_tmp / "report.docx"
            file_pdf = path_tmp / "report.pdf"
            _write_template(file_docx)
            captured_options: list[DocxToPdfOptions] = []

            def fake_convert(options: DocxToPdfOptions) -> DocxToPdfResult:
                captured_options.append(options)
                return DocxToPdfResult(file_pdf=options.file_out_pdf)

            with mock.patch(
                "docxrender.renderer.convert_docx_to_pdf",
                side_effect=fake_convert,
            ):
                result = (
                    DocxRenderer(file_docx=file_docx)
                    .with_pdf_conversion(
                        exe_libreoffice=Path("/usr/bin/libreoffice"),
                        dir_user_profile=path_tmp / "lo-profile",
                        file_out_pdf=file_pdf,
                    )
                    .write_pdf()
                )

            assert result.file_pdf == file_pdf
            assert captured_options[0].file_in_docx == file_docx

    def test_docx_renderer_pdf_options_use_field_marker_state(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_docx = path_tmp / "report.docx"
            options = (
                DocxRenderer(file_docx=file_docx)
                .with_field_update_markers(should_update_fields=False)
                .with_pdf_conversion(
                    exe_libreoffice=Path("/usr/bin/libreoffice"),
                    dir_user_profile=path_tmp / "lo-profile",
                    file_out_pdf=path_tmp / "report.pdf",
                )
                .pdf_options
            )

            assert options is not None
            assert options.should_update_fields is False
            assert options.should_freeze_fields is False

    def test_docx_renderer_does_not_expose_convert_docx_to_pdf(self) -> None:
        assert not hasattr(DocxRenderer(), "convert_docx_to_pdf")

    def test_parse_markdown_blocks_supports_commonmarkish_features(self) -> None:
        blocks = parse_markdown_blocks(
            "# Heading **Bold**\n\n"
            "Body with **bold**, `code`, and [link](https://example.com).  \n"
            "Body line 2.\n\n"
            "- Bullet one\n"
            "- Bullet **two**\n\n"
            "1. Ordered one\n"
            "2. Ordered two\n\n"
            "| A | B |\n"
            "| --- | :---: |\n"
            "| 1 | 2 |\n\n"
            "![Caption **Bold**](image.png){width=76%}\n"
        )

        assert blocks[0] == MarkdownHeading(
            level=1,
            text=(
                MarkdownTextSegment("Heading "),
                MarkdownTextSegment("Bold", is_bold=True),
            ),
        )
        assert blocks[1] == MarkdownParagraph(
            text=(
                MarkdownTextSegment("Body with "),
                MarkdownTextSegment("bold", is_bold=True),
                MarkdownTextSegment(", code, and link.\nBody line 2."),
            )
        )
        assert blocks[2] == MarkdownUnorderedList(
            items=(
                (MarkdownTextSegment("Bullet one"),),
                (
                    MarkdownTextSegment("Bullet "),
                    MarkdownTextSegment("two", is_bold=True),
                ),
            )
        )
        assert blocks[3] == MarkdownOrderedList(
            items=(
                (MarkdownTextSegment("Ordered one"),),
                (MarkdownTextSegment("Ordered two"),),
            )
        )
        assert blocks[4] == MarkdownTable(
            rows=(
                (
                    (MarkdownTextSegment("A"),),
                    (MarkdownTextSegment("B"),),
                ),
                (
                    (MarkdownTextSegment("1"),),
                    (MarkdownTextSegment("2"),),
                ),
            )
        )
        assert blocks[5] == MarkdownImage(
            caption=(
                MarkdownTextSegment("Caption "),
                MarkdownTextSegment("Bold", is_bold=True),
            ),
            path="image.png",
            width_pct=76.0,
        )

    def test_parse_markdown_blocks_can_disable_inline_markdown_rules(self) -> None:
        blocks = parse_markdown_blocks(
            "Body with **bold**, `code`, and [link](https://example.com).\n\n"
            "![Caption](image.png){width=76%}\n",
            options=DocxMarkdownOptions(
                should_parse_inline_bold=False,
                should_parse_inline_code=False,
                should_parse_links_as_text=False,
                should_parse_image_width_attr=False,
                default_image_width_pct=55.0,
            ),
        )

        assert blocks[0] == MarkdownParagraph(
            text=(
                MarkdownTextSegment(
                    "Body with **bold**, `code`, and [link](https://example.com)."
                ),
            )
        )
        assert blocks[1] == MarkdownImage(
            caption=(MarkdownTextSegment("Caption"),),
            path="image.png",
            width_pct=55.0,
        )

    def test_docx_renderer_write_pdf_writes_docx_before_pdf(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_docx = path_tmp / "report.docx"
            file_pdf = path_tmp / "report.pdf"
            _write_template(file_template)
            captured_options: list[DocxToPdfOptions] = []

            def fake_convert(options: DocxToPdfOptions) -> DocxToPdfResult:
                captured_options.append(options)
                return DocxToPdfResult(file_pdf=options.file_out_pdf)

            with mock.patch(
                "docxrender.renderer.convert_docx_to_pdf",
                side_effect=fake_convert,
            ):
                renderer = DocxRenderer().with_pdf_conversion(
                    exe_libreoffice=Path("/usr/bin/libreoffice"),
                    dir_user_profile=path_tmp / "lo-profile",
                    file_out_pdf=file_pdf,
                )
                result = renderer.write_pdf(
                    file_template=file_template,
                    file_out_docx=file_docx,
                    context={"report_title": "PDF"},
                    markdown_body="Body.",
                    dir_base=path_tmp,
                )

            assert result.file_pdf == file_pdf
            assert renderer.file_docx is None
            assert captured_options[0].file_in_docx == file_docx

    def test_write_docx_creates_minimal_document(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_image = path_tmp / "image.png"
            file_out_docx = path_tmp / "report.docx"
            _write_template(file_template)
            _write_png(file_image)

            options_docx = DocxWriteOptions(
                file_template=file_template,
                file_out_docx=file_out_docx,
                context={"report_title": "Contract Report"},
                markdown_body=(
                    "# Heading **Bold**\n\n"
                    "Body first line with **bold** and `code`.  \n"
                    "Body second line with [link](https://example.com).\n\n"
                    "注：Note text.\n\n"
                    "- Bullet item\n\n"
                    "1. First item\n"
                    "2. Second item\n\n"
                    "| A | B |\n"
                    "| --- | --- |\n"
                    "| 1 | 2 |\n\n"
                    "![Example image](image.png){width=76%}\n"
                ),
                dir_base=path_tmp,
                style=create_docx_style(),
            )

            result = write_docx(options_docx)
            document = Document(str(result.file_docx))
            paragraphs = document.paragraphs
            texts = [paragraph.text for paragraph in paragraphs]

            assert result.file_docx == file_out_docx
            assert "Contract Report" in texts
            assert "Heading Bold" in texts
            assert "__REPORT_BODY_ANCHOR__" not in texts
            assert (
                "Body first line with bold and code.\n"
                "Body second line with link."
            ) in texts
            assert "注：Note text." in texts
            assert "Bullet item" in texts
            assert "First item" in texts
            assert "Second item" in texts
            assert "Example image" in texts
            assert document.tables[0].cell(0, 0).text == "A"
            assert document.tables[0].cell(1, 1).text == "2"
            assert len(document.inline_shapes) == 1
            paragraph_by_text = {paragraph.text: paragraph for paragraph in paragraphs}
            assert (
                _run_font_size_pt(_first_text_run(paragraph_by_text["Heading Bold"]))
                == 16.0
            )
            assert (
                _run_font_size_pt(
                    _first_text_run(
                        paragraph_by_text[
                            "Body first line with bold and code.\n"
                            "Body second line with link."
                        ]
                    )
                )
                == 12.0
            )
            assert any(run.bold for run in paragraph_by_text["Heading Bold"].runs)
            assert any(
                run.bold
                for run in paragraph_by_text[
                    "Body first line with bold and code.\nBody second line with link."
                ].runs
                if run.text == "bold"
            )
            assert (
                _run_font_size_pt(_first_text_run(paragraph_by_text["注：Note text."]))
                == 10.5
            )
            assert (
                _run_font_size_pt(_first_text_run(paragraph_by_text["Example image"]))
                == 10.5
            )
            shape_image = cast(Any, document.inline_shapes[0])
            width_image = shape_image.width.inches
            assert width_image == pytest.approx(
                4.56,
                abs=0.02,
            )
            paragraph_caption = paragraph_by_text["Example image"]
            idx_caption = paragraphs.index(paragraph_caption)
            assert paragraphs[idx_caption - 1].alignment == 1
            table_markdown = document.tables[0]
            table_xml = cast(Any, table_markdown)._tbl.xml
            assert table_markdown.autofit is True
            assert '<w:tblLayout w:type="autofit"/>' in table_xml
            assert "<w:tblGrid>" in table_xml
            assert "<w:gridCol w:w=" not in table_xml
            assert "<w:tcW " not in table_xml
            assert 'w:val="single"' in cast(Any, table_markdown.cell(0, 0))._tc.xml

    def test_write_docx_body_render_policy_controls_structural_rendering(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_image = path_tmp / "image.png"
            file_out_docx = path_tmp / "report.docx"
            _write_template(file_template)
            _write_png(file_image)

            write_docx(
                DocxWriteOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={"report_title": "Contract Report"},
                    markdown_body=(
                        "# Heading\n\n"
                        "## Child\n\n"
                        "# 9. Existing\n\n"
                        "1. First\n"
                        "2. Second\n\n"
                        "- Bullet\n\n"
                        "| A | B |\n"
                        "| --- | --- |\n"
                        "| 1 | 2 |\n"
                        "| 3 | 4 |\n\n"
                        "![Caption](image.png){width=76%}\n"
                    ),
                    dir_base=path_tmp,
                    style=create_docx_style(),
                    body_render_policy=DocxBodyRenderPolicy(
                        should_number_headings=True,
                        rule_ordered_list="plain_text",
                        rule_unordered_list="plain_text",
                        should_stripe_table_rows=True,
                    ),
                )
            )

            document = Document(str(file_out_docx))
            paragraphs = document.paragraphs
            texts = [paragraph.text for paragraph in paragraphs]
            assert "1. Heading" in texts
            assert "1.1 Child" in texts
            assert "9. Existing" in texts
            assert "1. First" in texts
            assert "2. Second" in texts
            assert "- Bullet" in texts
            table_markdown = document.tables[0]
            body_shading = (
                cast(Any, table_markdown.rows[1].cells[0])
                ._tc.tcPr.first_child_found_in("w:shd")
            )
            assert body_shading is not None
            assert body_shading.get(qn("w:fill")) == "D9D9D9"
            assert cast(Any, table_markdown.cell(1, 0)).vertical_alignment is not None
            assert (
                cast(
                    Any,
                    table_markdown.cell(1, 0).paragraphs[0].paragraph_format,
                ).line_spacing
                == 1.5
            )
            paragraph_caption = next(
                paragraph
                for paragraph in paragraphs
                if paragraph.text == "Caption"
            )
            idx_caption = paragraphs.index(paragraph_caption)
            assert paragraphs[idx_caption - 1].alignment == 1
            assert paragraph_caption.alignment == 1
            paragraph_caption_format = cast(Any, paragraph_caption.paragraph_format)
            assert paragraph_caption_format.first_line_indent is None
            assert cast(Any, paragraph_caption.paragraph_format).line_spacing == 1.2

    def test_write_docx_template_renders_generic_docxtpl_context(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "template-rendered.docx"
            _write_template_with_anchor_text(file_template, "{{ body_anchor }}")

            result = write_docx_template(
                DocxTemplateRenderOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={"report_title": "Template Render"},
                    context_defaults={"body_anchor": "AUTO"},
                )
            )

            texts = [
                paragraph.text
                for paragraph in Document(str(result.file_docx)).paragraphs
            ]
            assert result.file_docx == file_out_docx
            assert "Template Render" in texts
            assert "AUTO" in texts

    def test_write_docx_template_can_use_defaults_win_conflict_policy(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "template-rendered.docx"
            _write_template_with_anchor_text(file_template, "{{ body_anchor }}")

            result = write_docx_template(
                DocxTemplateRenderOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={
                        "report_title": "Template Render",
                        "body_anchor": "EXPLICIT",
                    },
                    context_defaults={"body_anchor": "DEFAULT"},
                    context_policy=DocxTemplateContextPolicy(
                        rule_conflict="defaults_win"
                    ),
                )
            )

            texts = [
                paragraph.text
                for paragraph in Document(str(result.file_docx)).paragraphs
            ]
            assert result.file_docx == file_out_docx
            assert "DEFAULT" in texts
            assert "EXPLICIT" not in texts

    def test_write_docx_template_materializes_inline_images(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "template-rendered.docx"
            file_image = path_tmp / "cover.png"
            _write_template_with_anchor_text(file_template, "{{ cover_image }}")
            _write_png(file_image)

            result = write_docx_template(
                DocxTemplateRenderOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={"report_title": "Template Render"},
                    inline_images={
                        "cover_image": DocxTemplateImageSpec(
                            file_image=file_image,
                            width_mm=60,
                        )
                    },
                )
            )

            document = Document(str(result.file_docx))
            texts = [paragraph.text for paragraph in document.paragraphs]
            assert result.file_docx == file_out_docx
            assert "Template Render" in texts
            assert len(document.inline_shapes) == 1

    def test_write_docx_template_does_not_override_explicit_context(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "template-rendered.docx"
            _write_template_with_anchor_text(file_template, "{{ body_anchor }}")

            result = write_docx_template(
                DocxTemplateRenderOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={
                        "report_title": "Template Render",
                        "body_anchor": "EXPLICIT",
                    },
                    context_defaults={"body_anchor": "AUTO"},
                )
            )

            texts = [
                paragraph.text
                for paragraph in Document(str(result.file_docx)).paragraphs
            ]
            assert result.file_docx == file_out_docx
            assert "EXPLICIT" in texts
            assert "AUTO" not in texts

    def test_write_docx_template_can_require_context_keys(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "template-rendered.docx"
            _write_template_with_anchor_text(file_template, "{{ body_anchor }}")

            with pytest.raises(ValueError, match="Missing required template context"):
                write_docx_template(
                    DocxTemplateRenderOptions(
                        file_template=file_template,
                        file_out_docx=file_out_docx,
                        context={"report_title": "Template Render"},
                        context_defaults={},
                        context_policy=DocxTemplateContextPolicy(
                            required_keys=("report_title", "body_anchor")
                        ),
                    )
                )

    def test_write_docx_uses_template_context_policy_for_body_anchor_default(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template_with_anchor_text(file_template, "{{ body_anchor }}")

            write_docx(
                DocxWriteOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={
                        "report_title": "Conflict",
                        "body_anchor": "EXPLICIT",
                    },
                    markdown_body="Inserted body.",
                    dir_base=path_tmp,
                    style=create_docx_style(),
                    body_anchor=DocxBodyAnchorOptions(
                        anchor_token="DEFAULT",
                        rule_missing="raise",
                    ),
                    template_context_policy=DocxTemplateContextPolicy(
                        rule_conflict="defaults_win"
                    ),
                )
            )

            texts = [
                paragraph.text
                for paragraph in Document(str(file_out_docx)).paragraphs
            ]
            assert "Inserted body." in texts
            assert "DEFAULT" not in texts

    def test_write_docx_body_anchor_contains_match(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template_with_anchor_text(
                file_template,
                "before {{ body_anchor }} after",
            )

            write_docx(
                DocxWriteOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={"report_title": "Contains"},
                    markdown_body="Inserted body.",
                    dir_base=path_tmp,
                    style=create_docx_style(),
                    body_anchor=DocxBodyAnchorOptions(rule_match="contains"),
                )
            )

            texts = [
                paragraph.text for paragraph in Document(str(file_out_docx)).paragraphs
            ]
            assert "Inserted body." in texts
            assert "before __REPORT_BODY_ANCHOR__ after" not in texts

    def test_write_docx_does_not_overwrite_explicit_body_anchor_context(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template_with_anchor_text(file_template, "{{ body_anchor }}")

            write_docx(
                DocxWriteOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={"report_title": "Explicit", "body_anchor": "CUSTOM"},
                    markdown_body="Inserted body.",
                    dir_base=path_tmp,
                    style=create_docx_style(),
                    body_anchor=DocxBodyAnchorOptions(
                        anchor_token="CUSTOM",
                        rule_missing="raise",
                    ),
                )
            )

            texts = [
                paragraph.text for paragraph in Document(str(file_out_docx)).paragraphs
            ]
            assert "Inserted body." in texts
            assert "CUSTOM" not in texts

    def test_write_docx_body_anchor_missing_append(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template_with_anchor_text(file_template, "No anchor here")

            write_docx(
                DocxWriteOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={"report_title": "Append"},
                    markdown_body="Appended body.",
                    dir_base=path_tmp,
                    style=create_docx_style(),
                )
            )

            texts = [
                paragraph.text for paragraph in Document(str(file_out_docx)).paragraphs
            ]
            assert texts[-1] == "Appended body."

    def test_write_docx_body_anchor_missing_raise(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template_with_anchor_text(file_template, "No anchor here")

            with pytest.raises(ValueError, match="Could not locate body anchor"):
                write_docx(
                    DocxWriteOptions(
                        file_template=file_template,
                        file_out_docx=file_out_docx,
                        context={"report_title": "Raise"},
                        markdown_body="Body.",
                        dir_base=path_tmp,
                        style=create_docx_style(),
                        body_anchor=DocxBodyAnchorOptions(rule_missing="raise"),
                    )
                )

    def test_write_docx_body_anchor_multiple_matches_raise(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            document = Document()
            document.add_paragraph("{{ report_title }}")
            document.add_paragraph("{{ body_anchor }}")
            document.add_paragraph("{{ body_anchor }}")
            document.save(str(file_template))

            with pytest.raises(ValueError, match="multiple paragraphs"):
                write_docx(
                    DocxWriteOptions(
                        file_template=file_template,
                        file_out_docx=file_out_docx,
                        context={"report_title": "Multiple"},
                        markdown_body="Body.",
                        dir_base=path_tmp,
                        style=create_docx_style(),
                    )
                )

    def test_write_docx_body_anchor_preserves_section_properties(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            document = Document()
            document.add_paragraph("{{ report_title }}")
            document.add_paragraph("{{ body_anchor }}")
            document.add_section()
            document.save(str(file_template))

            write_docx(
                DocxWriteOptions(
                    file_template=file_template,
                    file_out_docx=file_out_docx,
                    context={"report_title": "Section"},
                    markdown_body="Inserted body.",
                    dir_base=path_tmp,
                    style=create_docx_style(),
                    body_anchor=DocxBodyAnchorOptions(rule_missing="raise"),
                )
            )

            texts = [
                paragraph.text for paragraph in Document(str(file_out_docx)).paragraphs
            ]
            assert "Inserted body." in texts
            assert "__REPORT_BODY_ANCHOR__" not in texts
            assert "<w:sectPr" in _read_docx_part(file_out_docx, "word/document.xml")

    def test_write_docx_without_field_refresh_does_not_import_uno(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template(file_template)

            def fail_on_uno_import(
                name: str,
                package: str | None = None,
            ) -> Any:
                if name == "uno":
                    raise AssertionError("Base write_docx must not import uno.")
                return import_module_original(name, package=package)

            import importlib

            import_module_original = importlib.import_module
            with mock.patch("importlib.import_module", side_effect=fail_on_uno_import):
                result = write_docx(
                    DocxWriteOptions(
                        file_template=file_template,
                        file_out_docx=file_out_docx,
                        context={"report_title": "No UNO"},
                        markdown_body="Body.",
                        dir_base=path_tmp,
                        style=create_docx_style(),
                    )
                )

            assert result.file_docx == file_out_docx
            assert file_out_docx.exists()

    def test_write_docx_field_refresh_overwrites_output_docx(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template(file_template)

            def fake_refresh_docx_with_uno(**kwargs: object) -> None:
                file_out = cast(Path, kwargs["file_out_docx"])
                file_out.write_bytes(b"refreshed")

            from docxrender import pdf_uno

            with mock.patch.object(
                pdf_uno,
                "refresh_docx_with_uno",
                side_effect=fake_refresh_docx_with_uno,
            ):
                result = write_docx(
                    DocxWriteOptions(
                        file_template=file_template,
                        file_out_docx=file_out_docx,
                        context={"report_title": "Refresh"},
                        markdown_body="Body.",
                        dir_base=path_tmp,
                        style=create_docx_style(),
                        field_refresh=DocxFieldRefreshOptions(
                            exe_libreoffice=Path("/usr/bin/libreoffice"),
                            dir_user_profile=path_tmp / "lo-profile",
                            poll_interval_seconds=0.0,
                            stable_checks=1,
                        ),
                    )
                )

            assert result.file_docx == file_out_docx
            assert file_out_docx.read_bytes() == b"refreshed"

    def test_write_docx_field_refresh_can_write_separate_output_docx(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            file_refreshed = path_tmp / "report-refreshed.docx"
            _write_template(file_template)

            def fake_refresh_docx_with_uno(**kwargs: object) -> None:
                file_in = cast(Path, kwargs["file_in_docx"])
                file_out = cast(Path, kwargs["file_out_docx"])
                assert file_in == file_out_docx
                file_out.write_bytes(b"refreshed separate")

            from docxrender import pdf_uno

            with mock.patch.object(
                pdf_uno,
                "refresh_docx_with_uno",
                side_effect=fake_refresh_docx_with_uno,
            ):
                result = write_docx(
                    DocxWriteOptions(
                        file_template=file_template,
                        file_out_docx=file_out_docx,
                        context={"report_title": "Refresh"},
                        markdown_body="Body.",
                        dir_base=path_tmp,
                        style=create_docx_style(),
                        field_refresh=DocxFieldRefreshOptions(
                            exe_libreoffice=Path("/usr/bin/libreoffice"),
                            dir_user_profile=path_tmp / "lo-profile",
                            file_out_docx_refreshed=file_refreshed,
                            poll_interval_seconds=0.0,
                            stable_checks=1,
                        ),
                    )
                )

            assert result.file_docx == file_out_docx
            assert file_out_docx.exists()
            assert file_refreshed.read_bytes() == b"refreshed separate"

    def test_write_docx_field_refresh_can_require_toc_and_freeze(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template(file_template)

            def fake_refresh_docx_with_uno(**kwargs: object) -> None:
                _write_minimal_field_docx(cast(Path, kwargs["file_out_docx"]))

            from docxrender import pdf_uno

            with mock.patch.object(
                pdf_uno,
                "refresh_docx_with_uno",
                side_effect=fake_refresh_docx_with_uno,
            ):
                write_docx(
                    DocxWriteOptions(
                        file_template=file_template,
                        file_out_docx=file_out_docx,
                        context={"report_title": "Refresh"},
                        markdown_body="Body.",
                        dir_base=path_tmp,
                        style=create_docx_style(),
                        field_refresh=DocxFieldRefreshOptions(
                            exe_libreoffice=Path("/usr/bin/libreoffice"),
                            dir_user_profile=path_tmp / "lo-profile",
                            should_require_toc=True,
                            should_freeze_fields=True,
                            poll_interval_seconds=0.0,
                            stable_checks=1,
                        ),
                    )
                )

            text_document = _read_docx_part(file_out_docx, "word/document.xml")
            assert "Rendered TOC" in text_document
            assert "fldChar" not in text_document

    def test_write_docx_field_refresh_reports_missing_required_toc(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_template = path_tmp / "template.docx"
            file_out_docx = path_tmp / "report.docx"
            _write_template(file_template)

            def fake_refresh_docx_with_uno(**kwargs: object) -> None:
                _write_minimal_field_docx_without_result(
                    cast(Path, kwargs["file_out_docx"])
                )

            from docxrender import pdf_uno

            with mock.patch.object(
                pdf_uno,
                "refresh_docx_with_uno",
                side_effect=fake_refresh_docx_with_uno,
            ):
                with pytest.raises(RuntimeError, match="TOC result"):
                    write_docx(
                        DocxWriteOptions(
                            file_template=file_template,
                            file_out_docx=file_out_docx,
                            context={"report_title": "Refresh"},
                            markdown_body="Body.",
                            dir_base=path_tmp,
                            style=create_docx_style(),
                            field_refresh=DocxFieldRefreshOptions(
                                exe_libreoffice=Path("/usr/bin/libreoffice"),
                                dir_user_profile=path_tmp / "lo-profile",
                                should_require_toc=True,
                                poll_interval_seconds=0.0,
                                stable_checks=1,
                            ),
                        )
                    )

    def test_create_libreoffice_listener_command_uses_isolated_profile(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_pdf_") as dir_tmp:
            command = create_libreoffice_listener_command(
                exe_libreoffice=Path("/usr/bin/libreoffice"),
                dir_user_profile=Path(dir_tmp) / "profile",
                port=23001,
            )

            assert command[0] == "/usr/bin/libreoffice"
            assert "--headless" in command
            assert "--accept=socket,host=127.0.0.1,port=23001;urp;" in command
            assert command[-1].startswith("-env:UserInstallation=file://")

    def test_convert_docx_to_pdf_stages_input_before_load(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_in_docx = path_tmp / "report.docx"
            file_in_docx.write_bytes(b"docx payload")
            file_source_lock = path_tmp / ".~lock.report.docx#"
            file_source_lock.write_text("lock", encoding="utf-8")
            file_out_pdf = path_tmp / "report.pdf"
            file_out_docx_refreshed = path_tmp / "report-refreshed.docx"
            file_listener_log = path_tmp / "listener.log"
            options_pdf = DocxToPdfOptions(
                exe_libreoffice=Path("/usr/bin/libreoffice"),
                file_in_docx=file_in_docx,
                file_out_pdf=file_out_pdf,
                dir_user_profile=path_tmp / "lo-profile",
                file_out_docx_refreshed=file_out_docx_refreshed,
                file_listener_log=file_listener_log,
                should_update_fields=False,
            )
            fake_process = FakeProcess()
            fake_doc = FakeDocument()
            fake_uno = types.SimpleNamespace(
                systemPathToFileUrl=create_fake_file_url
            )
            captured: dict[str, object] = {}

            def fake_load_document_or_raise(**kwargs: object) -> FakeDocument:
                captured.update(kwargs)
                file_staged = cast(Path, kwargs["file_in_docx_staged"])
                assert file_staged != file_in_docx
                assert file_staged.exists()
                assert file_staged.read_bytes() == file_in_docx.read_bytes()
                file_staged.write_bytes(b"refreshed staged payload")
                return fake_doc

            from docxrender import pdf_uno

            with (
                mock.patch.object(pdf_uno, "select_free_port", return_value=23001),
                mock.patch.object(
                    pdf_uno,
                    "create_libreoffice_listener_command",
                    return_value=["/usr/bin/libreoffice"],
                ),
                mock.patch.object(pdf_uno, "validate_libreoffice_executable"),
                mock.patch.object(
                    pdf_uno.subprocess,
                    "Popen",
                    return_value=fake_process,
                ),
                mock.patch.object(pdf_uno, "wait_for_listener"),
                mock.patch.object(pdf_uno, "import_uno_module", return_value=fake_uno),
                mock.patch.object(pdf_uno, "connect_desktop", return_value=object()),
                mock.patch.object(
                    pdf_uno,
                    "load_uno_document_or_raise",
                    side_effect=fake_load_document_or_raise,
                ),
                mock.patch.object(pdf_uno, "refresh_uno_document_fields"),
                mock.patch.object(
                    pdf_uno,
                    "create_property",
                    side_effect=create_fake_property,
                ),
            ):
                result = convert_docx_to_pdf(options_pdf)

            assert result.file_pdf == file_out_pdf
            assert result.file_docx_refreshed == file_out_docx_refreshed
            assert captured["file_in_docx_source"] == file_in_docx
            assert captured["file_source_lock"] == file_source_lock
            assert captured["file_listener_log"] == file_listener_log.resolve()
            assert fake_doc.stored
            assert fake_doc.closed
            assert len(fake_doc.store_url_calls) == 1
            assert file_out_docx_refreshed.read_bytes() == b"refreshed staged payload"
            assert fake_process.terminated

    def test_convert_docx_to_pdf_preserves_load_failure(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_contract_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_in_docx = path_tmp / "report.docx"
            file_in_docx.write_bytes(b"docx payload")
            options_pdf = DocxToPdfOptions(
                exe_libreoffice=Path("/usr/bin/libreoffice"),
                file_in_docx=file_in_docx,
                file_out_pdf=path_tmp / "report.pdf",
                dir_user_profile=path_tmp / "lo-profile",
                should_update_fields=False,
            )
            fake_process = FakeProcess()
            fake_uno = types.SimpleNamespace(
                systemPathToFileUrl=create_fake_file_url
            )

            from docxrender import pdf_uno

            with (
                mock.patch.object(pdf_uno, "select_free_port", return_value=23001),
                mock.patch.object(
                    pdf_uno,
                    "create_libreoffice_listener_command",
                    return_value=["/usr/bin/libreoffice"],
                ),
                mock.patch.object(pdf_uno, "validate_libreoffice_executable"),
                mock.patch.object(
                    pdf_uno.subprocess,
                    "Popen",
                    return_value=fake_process,
                ),
                mock.patch.object(pdf_uno, "wait_for_listener"),
                mock.patch.object(pdf_uno, "import_uno_module", return_value=fake_uno),
                mock.patch.object(pdf_uno, "connect_desktop", return_value=object()),
                mock.patch.object(
                    pdf_uno,
                    "load_uno_document_or_raise",
                    side_effect=RuntimeError("load failed"),
                ),
            ):
                with pytest.raises(RuntimeError, match="load failed"):
                    convert_docx_to_pdf(options_pdf)

            assert fake_process.terminated

    def test_pdf_load_failure_fields_include_staged_and_lock_info(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_pdf_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_source = path_tmp / "source.docx"
            file_source.write_bytes(b"source")
            file_staged = path_tmp / "staged.docx"
            file_staged.write_bytes(b"staged")
            file_listener_log = path_tmp / "listener.log"
            file_listener_log.write_text("listener tail", encoding="utf-8")
            file_source_lock = path_tmp / ".~lock.source.docx#"
            file_source_lock.write_text("lock", encoding="utf-8")

            fields = create_load_failure_fields(
                file_in_docx_source=file_source,
                file_in_docx_staged=file_staged,
                file_url="file:///tmp/staged.docx",
                exe_libreoffice=Path("/usr/bin/libreoffice"),
                dir_user_profile=path_tmp / "profile",
                process_listener=FakeProcess(),
                file_listener_log=file_listener_log,
                file_source_lock=file_source_lock,
                probe_ok=True,
                load_default_ok=False,
                load_hidden_only_ok=False,
            )
            text_error = "\n".join(fields)

            assert "error_code=libreoffice_uno_load_failed" in text_error
            assert "reason_code=staged_docx_import_failed" in text_error
            assert f"file_in_docx={file_source.resolve()}" in text_error
            assert f"file_in_docx_staged={file_staged.resolve()}" in text_error
            assert f"source_lock_file={file_source_lock.resolve()}" in text_error
            assert "probe_swriter_factory=ok" in text_error
            assert "load_staged_default_props=failed" in text_error
            assert "load_staged_hidden_only=failed" in text_error
            assert "listener_log_tail=listener tail" in text_error
            assert (
                "validate_libreoffice=libreoffice --headless --version" in text_error
            )
            assert (
                "install_debian_ubuntu=sudo apt install libreoffice python3-uno"
                in text_error
            )

    def test_pdf_uno_runtime_is_loaded_only_when_requested(self) -> None:
        with mock.patch("docxrender.pdf_uno.importlib.import_module") as import_module:
            import_module.side_effect = ImportError("missing uno")
            with pytest.raises(RuntimeError, match="libreoffice_uno_import_failed"):
                import_uno_module()
            import_module.assert_called_once_with("uno")

    def test_pdf_uno_import_failure_includes_install_guidance(self) -> None:
        with mock.patch("docxrender.pdf_uno.importlib.import_module") as import_module:
            import_module.side_effect = ImportError("missing uno")

            with pytest.raises(RuntimeError) as ctx:
                import_uno_module()

        text_error = str(ctx.value)
        assert "error_code=libreoffice_uno_import_failed" in text_error
        assert "validate_libreoffice=libreoffice --headless --version" in text_error
        assert 'validate_uno=python -c "import uno"' in text_error
        assert (
            "install_debian_ubuntu=sudo apt install libreoffice python3-uno"
            in text_error
        )
        assert "docs_libreoffice_parameters=" in text_error

    def test_convert_docx_to_pdf_reports_missing_libreoffice_executable(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_pdf_") as dir_tmp:
            path_tmp = Path(dir_tmp)
            file_in_docx = path_tmp / "report.docx"
            file_in_docx.write_bytes(b"docx payload")
            options_pdf = DocxToPdfOptions(
                exe_libreoffice=path_tmp / "missing-libreoffice",
                file_in_docx=file_in_docx,
                file_out_pdf=path_tmp / "report.pdf",
                dir_user_profile=path_tmp / "lo-profile",
                should_update_fields=False,
            )

            from docxrender import pdf_uno

            with (
                mock.patch.object(pdf_uno, "import_uno_module", return_value=object()),
                pytest.raises(FileNotFoundError) as ctx,
            ):
                convert_docx_to_pdf(options_pdf)

        text_error = str(ctx.value)
        assert "error_code=libreoffice_executable_missing" in text_error
        assert "missing-libreoffice" in text_error
        assert (
            "install_debian_ubuntu=sudo apt install libreoffice python3-uno"
            in text_error
        )

    def test_libreoffice_listener_start_failure_includes_guidance(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_pdf_") as dir_tmp:
            path_tmp = Path(dir_tmp)

            from docxrender import pdf_uno

            with (
                mock.patch.object(pdf_uno, "validate_libreoffice_executable"),
                mock.patch.object(
                    pdf_uno.subprocess,
                    "Popen",
                    side_effect=PermissionError("not executable"),
                ),
                pytest.raises(
                    RuntimeError,
                    match="libreoffice_listener_start_failed",
                ) as ctx,
            ):
                start_libreoffice_listener(
                    exe_libreoffice=Path("/usr/bin/libreoffice"),
                    dir_user_profile=path_tmp / "lo-profile",
                    port=23001,
                    stdout=None,
                    stderr=None,
                    file_listener_log=path_tmp / "listener.log",
                )

        text_error = str(ctx.value)
        assert "launch_error=PermissionError" in text_error
        assert "listener.log" in text_error
        assert "docs_libreoffice_api=" in text_error

    def test_libreoffice_listener_timeout_includes_guidance(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_pdf_") as dir_tmp:
            file_listener_log = Path(dir_tmp) / "listener.log"
            file_listener_log.write_text("cannot start", encoding="utf-8")

            from docxrender import pdf_uno

            with (
                mock.patch.object(pdf_uno, "LISTENER_START_TIMEOUT_SECONDS", 0.0),
                pytest.raises(
                    TimeoutError,
                    match="libreoffice_uno_listener_timeout",
                ) as ctx,
            ):
                wait_for_listener(23001, file_listener_log=file_listener_log)

        text_error = str(ctx.value)
        assert "listener_port=23001" in text_error
        assert "listener_log_tail=cannot start" in text_error
        assert "validate_libreoffice=libreoffice --headless --version" in text_error

    def test_docx_field_update_markers_are_written(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_fields_") as dir_tmp:
            file_docx = Path(dir_tmp) / "fields.docx"
            _write_minimal_field_docx(file_docx)

            write_docx_field_update_markers(file_docx)

            text_settings = _read_docx_part(file_docx, "word/settings.xml")
            text_document = _read_docx_part(file_docx, "word/document.xml")
            assert '<w:updateFields w:val="true"/>' in text_settings
            assert 'w:dirty="true"' in text_document

    def test_docx_field_freeze_preserves_result_text(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docxrender_fields_") as dir_tmp:
            file_docx = Path(dir_tmp) / "fields.docx"
            _write_minimal_field_docx(file_docx)
            write_docx_field_update_markers(file_docx)

            write_frozen_docx_fields(file_docx)

            text_settings = _read_docx_part(file_docx, "word/settings.xml")
            text_document = _read_docx_part(file_docx, "word/document.xml")
            assert "w:updateFields" not in text_settings
            assert "fldChar" not in text_document
            assert "instrText" not in text_document
            assert "w:dirty" not in text_document
            assert "Rendered TOC" in text_document

    def test_docxrender_does_not_import_product_repositories(self) -> None:
        product_module_prefixes = (
            "proteomics",
            "trait_association",
            "joint_proteome_phospho",
        )

        imported_product_modules = [
            name
            for name in sys.modules
            if name in product_module_prefixes
            or name.startswith(
                tuple(f"{prefix}." for prefix in product_module_prefixes)
            )
        ]

        assert imported_product_modules == []


def _write_template(file_template: Path) -> None:
    _write_template_with_anchor_text(file_template, "{{ body_anchor }}")


def _write_template_with_anchor_text(file_template: Path, anchor_text: str) -> None:
    document = Document()
    document.add_paragraph("{{ report_title }}")
    document.add_paragraph(anchor_text)
    document.save(str(file_template))


def _write_png(file_image: Path) -> None:
    file_image.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
            "AAAADUlEQVR42mP8z8BQDwAFgwJ/lv6OSwAAAABJRU5ErkJggg=="
        )
    )


def _write_minimal_field_docx(file_docx: Path) -> None:
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body><w:p>"
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        "<w:r><w:instrText>TOC</w:instrText></w:r>"
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        "<w:r><w:t>Rendered TOC</w:t></w:r>"
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        "</w:p></w:body></w:document>"
    )
    settings_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "</w:settings>"
    )
    with zipfile.ZipFile(file_docx, "w") as zip_out:
        zip_out.writestr("word/document.xml", document_xml)
        zip_out.writestr("word/settings.xml", settings_xml)


def _write_minimal_field_docx_without_result(file_docx: Path) -> None:
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body><w:p>"
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        "<w:r><w:instrText>TOC</w:instrText></w:r>"
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        "</w:p></w:body></w:document>"
    )
    settings_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "</w:settings>"
    )
    with zipfile.ZipFile(file_docx, "w") as zip_out:
        zip_out.writestr("word/document.xml", document_xml)
        zip_out.writestr("word/settings.xml", settings_xml)


def _read_docx_part(file_docx: Path, name: str) -> str:
    with zipfile.ZipFile(file_docx) as zip_in:
        return zip_in.read(name).decode("utf-8")


def _first_text_run(paragraph: Paragraph) -> Run:
    return next(run for run in paragraph.runs if run.text)


def _run_font_size_pt(run: Run) -> float:
    size = run.font.size
    if size is None:
        raise AssertionError("Expected text run to have an explicit font size.")
    return size.pt
